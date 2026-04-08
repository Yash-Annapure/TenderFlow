"""
Multi-format document parser.

Supported formats:
  .pdf   → PyMuPDF (text layer) with pdfplumber fallback for scanned/table-heavy PDFs
  .docx  → python-docx (paragraphs + tables)
  .md    → plain UTF-8
  .txt   → plain UTF-8
  .xlsx  → openpyxl (all sheets, row-by-row stringification)

Returns plain text for downstream chunking and embedding.
"""

import io
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


def parse_file(file_path: str, content: Optional[bytes] = None) -> str:
    """
    Parse a document to plain text.

    Args:
        file_path: Used to determine file type (by extension).
        content:   Raw file bytes. If None, the file is read from file_path.

    Returns:
        Extracted plain text string.

    Raises:
        ValueError: If the file type is unsupported or no text can be extracted.
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if content is None:
        with open(file_path, "rb") as fh:
            content = fh.read()

    if suffix == ".pdf":
        return _parse_pdf(content, filename=path.name)
    elif suffix == ".docx":
        return _parse_docx(content)
    elif suffix in (".md", ".txt"):
        return content.decode("utf-8", errors="replace").strip()
    elif suffix == ".xlsx":
        return _parse_xlsx(content)
    else:
        # Last-resort: attempt UTF-8 decode
        try:
            text = content.decode("utf-8", errors="replace").strip()
            if len(text) < 20:
                raise ValueError(f"Unsupported or empty file type: {suffix}")
            return text
        except Exception:
            raise ValueError(f"Unsupported file type: {suffix}")


# ── PDF ──────────────────────────────────────────────────────────────────────

def _parse_pdf(content: bytes, filename: str = "") -> str:
    import fitz  # PyMuPDF

    parts: list[str] = []
    with fitz.open(stream=content, filetype="pdf") as doc:
        for page in doc:
            parts.append(page.get_text("text"))

    text = "\n".join(parts).strip()

    # Fall back to pdfplumber when PyMuPDF yields too little
    # (common with scanned PDFs or heavily-formatted tables)
    if len(text) < 100:
        logger.info(f"PyMuPDF yielded sparse text for {filename}, falling back to pdfplumber")
        text = _parse_pdf_pdfplumber(content)

    return text


def _parse_pdf_pdfplumber(content: bytes) -> str:
    import pdfplumber

    parts: list[str] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            if page_text:
                parts.append(page_text)
            # Stringify tables as tab-separated rows
            for table in page.extract_tables():
                for row in table:
                    parts.append("\t".join(str(cell or "").strip() for cell in row))

    return "\n".join(parts).strip()


# ── DOCX ─────────────────────────────────────────────────────────────────────

def _parse_docx(content: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(content))
    parts: list[str] = []

    for para in doc.paragraphs:
        stripped = para.text.strip()
        if stripped:
            parts.append(stripped)

    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    return "\n\n".join(parts)


# ── XLSX ─────────────────────────────────────────────────────────────────────

def _parse_xlsx(content: bytes) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    parts: list[str] = []

    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        parts.append(f"=== Sheet: {sheet_name} ===")
        for row in sheet.iter_rows(values_only=True):
            row_text = "\t".join(str(cell) if cell is not None else "" for cell in row)
            if row_text.strip():
                parts.append(row_text)

    wb.close()
    return "\n".join(parts)
