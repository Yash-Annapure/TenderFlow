"""
DOCX output renderer.

Converts TenderState into a professionally formatted Word document:
  Cover page          — title, date, final readiness score + band
  Readiness Assessment — score table (Primary / Compliance / Robustness / Final)
                         + score justifications
                         + Action Items checklist (gap flags)
  Tender Sections     — one heading-2 per section, confidence indicator,
                         gap warning (if any), body text, sources footer

Content priority per section:
  finalised_content  (post-HITL Sonnet polish)
  > user_edits       (raw human edits, used if finalise didn't run)
  > draft_text       (original AI draft)

python-docx is used directly — no system-level PDF dependencies needed.
"""

import logging
import os
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from config.settings import settings

logger = logging.getLogger(__name__)

# ── Colour palette ─────────────────────────────────────────────────────────────
_BAND_COLOURS: dict[str, RGBColor] = {
    "EXCELLENT": RGBColor(0x27, 0xAE, 0x60),
    "STRONG": RGBColor(0x29, 0x80, 0xB9),
    "MODERATE": RGBColor(0xF3, 0x9C, 0x12),
    "WEAK": RGBColor(0xE7, 0x4C, 0x3C),
}
_GREY = RGBColor(0x7F, 0x8C, 0x8D)
_WARN_COLOUR = RGBColor(0xE7, 0x4C, 0x3C)

_CONFIDENCE_LABELS = {
    "HIGH": "[HIGH CONFIDENCE]",
    "MEDIUM": "[MEDIUM CONFIDENCE]",
    "LOW": "[LOW CONFIDENCE]",
}


_XML_ILLEGAL = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]')
_MARKDOWN_BOLD_ITALIC = re.compile(r'\*{1,3}([^*]+)\*{1,3}')
_MARKDOWN_HEADING = re.compile(r'^#{1,6}\s+', re.MULTILINE)
_TABLE_ROW = re.compile(r'^\s*\|.+\|\s*$')
_TABLE_SEP = re.compile(r'^\s*\|[\s\-:|]+\|\s*$')


def _sanitize(text: str) -> str:
    """Strip XML-illegal control characters and markdown syntax from DOCX text."""
    text = _XML_ILLEGAL.sub('', text).replace('\r\n', '\n').replace('\r', '\n')
    text = _MARKDOWN_BOLD_ITALIC.sub(r'\1', text)
    text = _MARKDOWN_HEADING.sub('', text)
    return text


def _sanitize_cell(text: str) -> str:
    """Sanitize a table cell — keep stars for bold detection before stripping."""
    return _XML_ILLEGAL.sub('', text).replace('\r\n', ' ').replace('\r', ' ').replace('\n', ' ')


def _is_table_row(line: str) -> bool:
    return bool(_TABLE_ROW.match(line))


def _is_separator_row(line: str) -> bool:
    return bool(_TABLE_SEP.match(line))


def _parse_row_cells(line: str) -> list[str]:
    """Split a pipe-delimited row into cell strings."""
    parts = line.strip().split('|')
    # Remove first and last (empty from leading/trailing |)
    return [_sanitize_cell(p.strip()) for p in parts[1:-1]]


def _is_bold_cell(text: str) -> bool:
    return text.startswith('**') and text.endswith('**')


def _unwrap_bold(text: str) -> str:
    return _MARKDOWN_BOLD_ITALIC.sub(r'\1', text)


def _add_docx_table(doc: Document, rows: list[list[str]]) -> None:
    """Render a parsed markdown table as a styled DOCX table."""
    if not rows:
        return

    navy = RGBColor(0x1D, 0x35, 0x57)
    white = RGBColor(0xFF, 0xFF, 0xFF)
    light_bg = RGBColor(0xF0, 0xF4, 0xF8)

    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"

    for ri, row_cells in enumerate(rows):
        tr = table.rows[ri]
        for ci in range(col_count):
            cell_text = row_cells[ci] if ci < len(row_cells) else ""
            is_bold = _is_bold_cell(cell_text)
            clean_text = _unwrap_bold(cell_text)

            cell = tr.cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(clean_text)
            run.font.size = Pt(10)

            if ri == 0:
                # Header row — navy background, white bold text
                run.bold = True
                run.font.color.rgb = white
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), '1D3557')
                tcPr.append(shd)
            elif ri % 2 == 0:
                # Even body rows — light blue-grey
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F0F4F8')
                tcPr.append(shd)

            if is_bold or ri == 0:
                run.bold = True

    doc.add_paragraph()


def _score_band(score: float) -> str:
    if score >= 90:
        return "EXCELLENT"
    if score >= 75:
        return "STRONG"
    if score >= 60:
        return "MODERATE"
    return "WEAK"


# ── Public entry point ─────────────────────────────────────────────────────────

def render_docx(state: dict, tender_id: str) -> str:
    """
    Render a complete tender response DOCX from the agent state dict.

    Args:
        state:     TenderState dict (serialisable subset — no LangGraph internals).
        tender_id: Used as the output filename.

    Returns:
        Absolute path to the saved .docx file.
    """
    os.makedirs(settings.outputs_dir, exist_ok=True)
    output_path = str(Path(settings.outputs_dir) / f"{tender_id}.docx")

    doc = Document()
    _configure_styles(doc)
    _add_cover(doc, state)
    doc.add_page_break()
    _add_readiness_assessment(doc, state)
    doc.add_page_break()

    for section in state.get("sections", []):
        _add_section(doc, section)

    doc.save(output_path)
    logger.info(f"[output] DOCX saved → {output_path}")
    return output_path


# ── Document sections ──────────────────────────────────────────────────────────

def _configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def _add_cover(doc: Document, state: dict) -> None:
    # Derive a clean tender title from the filename
    raw_filename = state.get("tender_filename", "Tender Response")
    tender_title = re.sub(r'\.(pdf|docx?|txt)$', '', raw_filename, flags=re.IGNORECASE)
    tender_title = re.sub(r'[_\-]+', ' ', tender_title).strip().title()

    # Navy background via a shaded table (1×1 cover block)
    cover_table = doc.add_table(rows=1, cols=1)
    cover_cell = cover_table.rows[0].cells[0]
    tcPr = cover_cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1D3557')
    tcPr.append(shd)

    # Category label (top right feel via separate paragraph)
    cp0 = cover_cell.paragraphs[0]
    cp0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r0 = cp0.add_run("Technical Tender")
    r0.italic = True
    r0.font.size = Pt(10)
    r0.font.color.rgb = RGBColor(0xA8, 0xC4, 0xE0)

    cp1 = cover_cell.add_paragraph()
    cp1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = cp1.add_run(tender_title)
    r1.bold = True
    r1.font.size = Pt(20)
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    cp2 = cover_cell.add_paragraph()
    r2 = cp2.add_run("Meridian Intelligence GmbH")
    r2.italic = True
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0xA8, 0xC4, 0xE0)

    doc.add_paragraph()

    # Meta row
    meta = doc.add_table(rows=1, cols=3)
    meta_cells = meta.rows[0].cells
    meta_cells[0].text = f"Generated by: TenderFlow AI Agent"
    meta_cells[1].text = f"Date: {datetime.now().strftime('%d %B %Y')}"
    final_score = state.get("final_score", 0.0)
    band = _score_band(final_score)
    meta_cells[2].text = f"Readiness: {final_score:.1f}/100 ({band})"
    for mc in meta_cells:
        for run in mc.paragraphs[0].runs:
            run.font.size = Pt(9)
            run.font.color.rgb = _GREY

    doc.add_paragraph()
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.add_run(
        "For internal review only · Do not submit without human review"
    ).italic = True


def _add_readiness_assessment(doc: Document, state: dict) -> None:
    doc.add_heading("Readiness Assessment", level=1)

    # Score summary table
    score_rows = [
        ("Primary Score (Track Record + Expertise + Methodology)", state.get("primary_score_total", 0.0)),
        ("Compliance Coverage", state.get("compliance_score", 0.0)),
        ("Robustness Index", state.get("robustness_score", 0.0)),
        ("Final Score", state.get("final_score", 0.0)),
    ]

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Shading Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Dimension"
    hdr[1].text = "Score"
    hdr[2].text = "Band"
    for hdr_cell in hdr:
        for run in hdr_cell.paragraphs[0].runs:
            run.bold = True

    for dimension, score in score_rows:
        row = table.add_row().cells
        row[0].text = dimension
        row[1].text = f"{score:.1f} / 100"
        band = _score_band(score)
        row[2].text = band

    doc.add_paragraph()

    # Module breakdown table (M1-M5)
    primary_scores = state.get("primary_scores", {})
    if primary_scores:
        doc.add_heading("Primary Score Breakdown", level=2)
        _MODULE_LABELS = {
            "M1_track_record": "Track Record — past project similarity",
            "M2_expertise_depth": "Expertise Depth — domain & regulatory coverage",
            "M3_methodology_fit": "Methodology Fit — approach alignment",
            "M4_delivery_credibility": "Delivery Credibility — team CV match",
            "M5_pricing": "Pricing Proxy — budget competitiveness",
        }
        mod_table = doc.add_table(rows=1, cols=3)
        mod_table.style = "Light Shading Accent 1"
        mhdr = mod_table.rows[0].cells
        mhdr[0].text = "Module"
        mhdr[1].text = "Score"
        mhdr[2].text = "Band"
        for c in mhdr:
            for run in c.paragraphs[0].runs:
                run.bold = True
        for key, label in _MODULE_LABELS.items():
            score = primary_scores.get(key)
            if score is None:
                continue
            mrow = mod_table.add_row().cells
            mrow[0].text = label
            mrow[1].text = f"{score:.0f} / 100"
            mrow[2].text = _score_band(score)
        doc.add_paragraph()

    # Score justifications
    justifications = state.get("score_justifications", {})
    if justifications:
        doc.add_heading("Score Justifications", level=2)
        for module, justification in justifications.items():
            p = doc.add_paragraph()
            p.add_run(f"{module}:  ").bold = True
            p.add_run(_sanitize(str(justification)))

    # Action Items — gap flags + LOW-confidence sections
    doc.add_heading("Action Items", level=2)
    sections = state.get("sections", [])
    gap_sections   = [s for s in sections if s.get("gap_flag")]
    low_no_gap     = [s for s in sections if s.get("confidence") == "LOW" and not s.get("gap_flag")]

    if gap_sections or low_no_gap:
        for s in gap_sections:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"[{s['section_name']}]  ")
            run.bold = True
            run.font.color.rgb = _WARN_COLOUR
            p.add_run(s["gap_flag"])
        for s in low_no_gap:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"[{s['section_name']}]  ")
            run.bold = True
            run.font.color.rgb = _WARN_COLOUR
            p.add_run(
                "Low-confidence draft — KB content found via fallback retrieval only. "
                "Review and supplement with project-specific evidence."
            )
    else:
        doc.add_paragraph("No critical knowledge-base gaps identified for this tender.")


def _render_content(doc: Document, text: str) -> None:
    """
    Render section body text into the document.
    Handles:
      - Markdown pipe tables  → styled DOCX tables
      - ### / ## / # headings → Heading 3 / 4 / normal bold
      - Bullet / numbered lists
      - Prose paragraphs
    """
    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── Skip empty lines ────────────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── Markdown table block ────────────────────────────────────────────
        if _is_table_row(stripped):
            table_lines = []
            while i < len(lines) and _is_table_row(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1
            # Parse: first row = header, skip separator rows, rest = body
            parsed_rows = []
            for tl in table_lines:
                if _is_separator_row(tl):
                    continue
                cells = _parse_row_cells(tl)
                if cells:
                    parsed_rows.append(cells)
            _add_docx_table(doc, parsed_rows)
            continue

        # ── Sub-headings (e.g. "3.1 Identification") ───────────────────────
        h3 = stripped.startswith('### ')
        h2 = stripped.startswith('## ')
        h1 = stripped.startswith('# ')
        if h3 or h2 or h1:
            prefix_len = 4 if h3 else 3 if h2 else 2
            heading_text = _sanitize(stripped[prefix_len:])
            p = doc.add_paragraph()
            run = p.add_run(heading_text)
            run.bold = True
            run.font.size = Pt(11)
            i += 1
            continue

        # ── Numbered sub-section headers (e.g. "3.1 Identification") ───────
        # Detect lines like "3.1 Something" that look like sub-headings
        import re as _re
        subsection = _re.match(r'^(\d+\.\d+)\s+(.+)$', stripped)
        if subsection and len(stripped) < 80:
            p = doc.add_paragraph()
            run = p.add_run(f"{subsection.group(1)} {_sanitize(subsection.group(2))}")
            run.bold = True
            run.font.size = Pt(11)
            i += 1
            continue

        # ── Bullet list ─────────────────────────────────────────────────────
        if stripped.startswith(('- ', '• ', '* ')):
            while i < len(lines) and lines[i].strip().startswith(('- ', '• ', '* ')):
                item = _sanitize(lines[i].strip()[2:])
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(item)
                i += 1
            continue

        # ── Numbered list ───────────────────────────────────────────────────
        num_match = _re.match(r'^\d+[.)]\s+', stripped)
        if num_match:
            while i < len(lines):
                nl = lines[i].strip()
                if not _re.match(r'^\d+[.)]\s+', nl):
                    break
                item = _sanitize(_re.sub(r'^\d+[.)]\s+', '', nl))
                p = doc.add_paragraph(style='List Number')
                p.add_run(item)
                i += 1
            continue

        # ── Regular paragraph ───────────────────────────────────────────────
        para_lines = []
        while i < len(lines):
            pl = lines[i].strip()
            if not pl:
                break
            if _is_table_row(pl):
                break
            if pl.startswith(('# ', '## ', '### ', '- ', '• ', '* ')):
                break
            if _re.match(r'^\d+[.)]\s+', pl):
                break
            para_lines.append(_sanitize(pl))
            i += 1
        if para_lines:
            p = doc.add_paragraph()
            # Handle bold/italic within paragraph lines
            full_text = ' '.join(para_lines)
            # Split on **bold** markers
            parts = _re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', full_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    p.add_run(part)


def _add_section(doc: Document, section: dict) -> None:
    section_name = section.get("section_name", "Untitled Section")
    confidence = section.get("confidence", "MEDIUM")
    gap_flag = section.get("gap_flag")

    # Each section starts on its own page
    doc.add_page_break()

    # Heading with inline confidence indicator
    heading = doc.add_heading(level=2)
    heading.add_run(_sanitize(section_name))
    conf_run = heading.add_run(f"   {_CONFIDENCE_LABELS.get(confidence, confidence)}")
    conf_run.font.size = Pt(9)
    conf_run.bold = False
    conf_run.font.color.rgb = _GREY

    # Gap warning
    if gap_flag:
        p = doc.add_paragraph()
        warn = p.add_run(f"Knowledge-base gap: {_sanitize(gap_flag)}")
        warn.italic = True
        warn.font.color.rgb = _WARN_COLOUR

    # Body content — priority: finalised > user_edits > draft_text
    raw_content: str = (
        section.get("finalised_content")
        or section.get("user_edits")
        or section.get("draft_text")
        or ""
    )
    raw_content = _XML_ILLEGAL.sub('', raw_content).replace('\r\n', '\n').replace('\r', '\n')

    if raw_content.strip():
        _render_content(doc, raw_content)
    else:
        p = doc.add_paragraph()
        p.add_run("[No content generated for this section]").italic = True

    # Sources footer
    sources = section.get("sources_used", [])
    if sources:
        p = doc.add_paragraph()
        run = p.add_run("Sources: " + " · ".join(_sanitize(s) for s in sources))
        run.font.size = Pt(9)
        run.italic = True
        run.font.color.rgb = _GREY

    # Visual separator
    doc.add_paragraph()
